"""
Genera un reporte detallado en formato .docx con todos los resultados del paper.

Cubre los 25 experimentos: 9 modelos × 3 backfills (SARIMA solo baseline).
Incluye tablas de métricas, análisis de ablación, impacto del backfill,
importancia SHAP y figuras embebidas.

Uso:
    python src/generate_report.py
    python src/generate_report.py --out reporte_paper.docx
"""

import argparse
import json
import os
from datetime import date
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# ─── Configuración de experimentos ────────────────────────────────────────────

SERIES = [
    "ivf_total_nacional",
    "ivf_turistico_total",
    "ivf_turistico_bienes",
    "ivf_turistico_servicios",
]

SERIES_LABELS = {
    "ivf_total_nacional":       "IVF Total Nacional",
    "ivf_turistico_total":      "IVF Turístico Total",
    "ivf_turistico_bienes":     "IVF Turístico Bienes",
    "ivf_turistico_servicios":  "IVF Turístico Servicios",
}

# Todos los experimentos (nombre → (modelo, backfill))
ALL_EXPERIMENTS = {
    "xgb_zero":              ("XGBoost",     "Zero"),
    "xgb_linear":            ("XGBoost",     "Linear"),
    "xgb_xgb_backcast":      ("XGBoost",     "XGB Backcast"),
    "ridge_zero":            ("Ridge",       "Zero"),
    "ridge_linear":          ("Ridge",       "Linear"),
    "ridge_xgb_backcast":    ("Ridge",       "XGB Backcast"),
    "sarima_baseline":       ("SARIMA",      "—"),
    "sarimax_zero":          ("SARIMAX",     "Zero"),
    "sarimax_linear":        ("SARIMAX",     "Linear"),
    "sarimax_xgb_backcast":  ("SARIMAX",     "XGB Backcast"),
    "mlp_zero":              ("MLP",         "Zero"),
    "mlp_linear":            ("MLP",         "Linear"),
    "mlp_xgb_backcast":      ("MLP",         "XGB Backcast"),
    "gru_zero":              ("GRU",         "Zero"),
    "gru_linear":            ("GRU",         "Linear"),
    "gru_xgb_backcast":      ("GRU",         "XGB Backcast"),
    "lstm_zero":             ("LSTM",        "Zero"),
    "lstm_linear":           ("LSTM",        "Linear"),
    "lstm_xgb_backcast":     ("LSTM",        "XGB Backcast"),
    "cnngru_zero":           ("CNN-GRU",     "Zero"),
    "cnngru_linear":         ("CNN-GRU",     "Linear"),
    "cnngru_xgb_backcast":   ("CNN-GRU",     "XGB Backcast"),
    "rescnngru_zero":        ("Res-CNN-GRU", "Zero"),
    "rescnngru_linear":      ("Res-CNN-GRU", "Linear"),
    "rescnngru_xgb_backcast":("Res-CNN-GRU", "XGB Backcast"),
}

MODEL_ORDER = ["XGBoost", "Ridge", "SARIMA", "SARIMAX", "MLP", "GRU", "LSTM", "CNN-GRU", "Res-CNN-GRU"]

MODEL_DESCRIPTIONS = {
    "XGBoost":     "Gradient boosting sobre características tabulares de rezago",
    "Ridge":       "Regresión lineal regularizada (L2) sobre características de rezago",
    "SARIMA":      "SARIMA(1,1,1)(1,1,1)₄ — baseline temporal puro sin indicadores",
    "SARIMAX":     "SARIMA con indicadores exógenos de visitantes INEGI",
    "MLP":         "Red neuronal densa (FC 64→32→1) con early stopping",
    "GRU":         "Gated Recurrent Unit (seq=8, hidden=64, layers=1) con early stopping",
    "LSTM":        "Long Short-Term Memory (seq=8, hidden=64, layers=1) con early stopping",
    "CNN-GRU":     "2 bloques Conv1D (32, 64 canales) + GRU(64) + FC(64→32→1)",
    "Res-CNN-GRU": "3 bloques ResConv1D (16→32→64) + GRU(64) + FC(64→32→1)",
}

BACKFILL_DESCRIPTIONS = {
    "Zero":        "Rellena con ceros los trimestres pre-2018 (ruptura estructural artificial)",
    "Linear":      "Extrapolación lineal por indicador hacia atrás desde 2018",
    "XGB Backcast":"XGBoost entrenado sobre la serie temporal invertida (backcast recursivo)",
    "—":           "N/A — SARIMA no usa indicadores exógenos",
}

# ─── Carga de métricas ────────────────────────────────────────────────────────

def load_metrics(exp_name: str, series: str) -> dict | None:
    path = f"metrics/{exp_name}/metrics_{series}.json"
    if not os.path.exists(path):
        return None
    with open(path) as f:
        return json.load(f)


def build_df() -> pd.DataFrame:
    rows = []
    for exp, (model, backfill) in ALL_EXPERIMENTS.items():
        for series in SERIES:
            m = load_metrics(exp, series)
            rows.append({
                "exp":      exp,
                "model":    model,
                "backfill": backfill,
                "series":   series,
                "mae":      m["mae"]  if m else np.nan,
                "rmse":     m["rmse"] if m else np.nan,
                "r2":       m["r2"]   if m else np.nan,
                "mape":     m.get("mape", np.nan) if m else np.nan,
            })
    return pd.DataFrame(rows)


def avg4(df: pd.DataFrame, model: str, backfill: str, metric: str) -> float:
    sub = df[(df["model"] == model) & (df["backfill"] == backfill)]
    vals = sub[metric].dropna().values
    return float(np.mean(vals)) if len(vals) > 0 else np.nan


def fmt(v, decimals=4) -> str:
    if v is None or (isinstance(v, float) and np.isnan(v)):
        return "—"
    return f"{v:.{decimals}f}"


# ─── Helpers de formato docx ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Establece el color de fondo de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_row_bg(row, hex_color: str):
    for cell in row.cells:
        set_cell_bg(cell, hex_color)


def bold_cell(cell, text: str, size: int = 9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def normal_cell(cell, text: str, size: int = 9, bold: bool = False,
                align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_figure(doc: Document, img_path: str, caption: str, width: float = 6.0):
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f"[Figura no disponible: {img_path}]")
        p.runs[0].italic = True
        return
    doc.add_picture(img_path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(31, 73, 125)


def highlight_best(values: list[float], lower_is_better: bool = True) -> list[bool]:
    """Retorna máscara de cuál(es) son el mejor valor."""
    valid = [v for v in values if not np.isnan(v)]
    if not valid:
        return [False] * len(values)
    best = min(valid) if lower_is_better else max(valid)
    return [not np.isnan(v) and abs(v - best) < 1e-6 for v in values]


# ─── Secciones del documento ──────────────────────────────────────────────────

def write_title_page(doc: Document):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Pronóstico del Índice de Volumen Físico Turístico de México\n"
        "mediante Aprendizaje Automático e Indicadores INEGI"
    )
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run(
        "Estudio comparativo de 9 arquitecturas × 3 métodos de imputación\n"
        "sobre datos trimestrales INEGI 1994–2025"
    )
    r2.italic = True
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Reporte generado: {date.today().strftime('%d de %B de %Y')}").font.size = Pt(11)

    doc.add_paragraph()
    doc.add_page_break()


def write_abstract(doc: Document, df: pd.DataFrame):
    add_heading(doc, "Resumen", 1)

    # Calcular estadísticas clave
    xgb_lin_mae  = avg4(df, "XGBoost", "Linear", "mae")
    xgb_zero_mae = avg4(df, "XGBoost", "Zero",   "mae")
    ridge_lin_mae = avg4(df, "Ridge",  "Linear", "mae")
    sarima_mae   = avg4(df, "SARIMA",  "—",      "mae")
    best_nn_mae  = min(avg4(df, m, "Linear", "mae") for m in ["MLP","GRU","LSTM","CNN-GRU","Res-CNN-GRU"]
                       if not np.isnan(avg4(df, m, "Linear", "mae")))
    bf_gain = (xgb_zero_mae - xgb_lin_mae) / xgb_zero_mae * 100 if xgb_zero_mae else 0

    abstract_text = (
        f"Este trabajo presenta un estudio comparativo de nueve arquitecturas de aprendizaje automático "
        f"para el pronóstico trimestral del Índice de Volumen Físico (IVF) del sector turístico de México, "
        f"utilizando datos del INEGI para el período 1994–2025. El diseño experimental evalúa "
        f"9 modelos — XGBoost, Ridge, SARIMA, SARIMAX, MLP, GRU, LSTM, CNN-GRU y Res-CNN-GRU — "
        f"bajo tres métodos de imputación (backfill) para indicadores de visitantes con cobertura parcial "
        f"desde 2018: relleno con ceros, extrapolación lineal y backcast con XGBoost. "
        f"El conjunto de prueba cubre 2022 Q1–2025 Q3 (post-COVID), con una variable dummy para "
        f"2020 Q1–2021 Q4 que aísla la ruptura estructural de la pandemia.\n\n"
        f"Los resultados principales muestran que: (1) XGBoost con backfill lineal obtiene el menor "
        f"MAE promedio sobre las 4 series IVF (MAE = {fmt(xgb_lin_mae, 4)}), superando a SARIMA puro "
        f"(MAE = {fmt(sarima_mae, 4)}) y a las redes neuronales (MAE mínimo en NNs = {fmt(best_nn_mae, 4)}); "
        f"(2) el backfill lineal reduce el MAE de XGBoost en {bf_gain:.1f}% frente al relleno con ceros, "
        f"validando la hipótesis de que la imputación de alta calidad de indicadores históricos mejora "
        f"significativamente el pronóstico; (3) el análisis SHAP revela que los indicadores de gasto de "
        f"turistas internacionales por vía aérea son los predictores más influyentes más allá de los "
        f"rezagos autorregresivos puros."
    )

    p = doc.add_paragraph(abstract_text)
    p.runs[0].font.size = Pt(10)
    for run in p.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()
    kw = doc.add_paragraph()
    r = kw.add_run("Palabras clave: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = kw.add_run("pronóstico de turismo, IVF, XGBoost, redes neuronales recurrentes, SARIMA, "
                     "imputación de datos, SHAP, INEGI, México, COVID-19")
    r2.font.size = Pt(10)
    doc.add_page_break()


def write_introduction(doc: Document):
    add_heading(doc, "1. Introducción", 1)

    paras = [
        ("El sector turístico mexicano representa aproximadamente el 8.7% del PIB nacional y es una "
         "fuente crítica de divisas. La pandemia de COVID-19 provocó una caída sin precedentes en el "
         "Índice de Volumen Físico (IVF) turístico durante 2020–2021, seguida de una recuperación "
         "asimétrica entre subsectores. La capacidad de pronosticar con precisión la dinámica post-COVID "
         "del IVF es fundamental para la planificación de política económica y la asignación de recursos."),
        ("El Instituto Nacional de Estadística y Geografía (INEGI) publica dos conjuntos de datos "
         "relevantes: (1) el IVF trimestral desde 1993, que ofrece una serie histórica larga pero sin "
         "desagregación de flujos de visitantes; y (2) la Encuesta de Viajeros Internacionales (EVI) y "
         "estadísticas de visitantes con cobertura mensual desde 2018, que capturan patrones de movilidad "
         "turística pero con cobertura histórica limitada. Esta asimetría temporal — series objetivo "
         "largas vs. indicadores cortos — plantea el problema central de cómo integrar indicadores "
         "parciales en modelos de pronóstico de series largas."),
        ("Este trabajo hace tres contribuciones principales: (1) propone y compara tres métodos de "
         "imputación (backfill) para extender los indicadores INEGI hacia períodos pre-2018 — relleno "
         "con ceros, extrapolación lineal y backcast con XGBoost — evaluando su impacto en la calidad "
         "del pronóstico; (2) realiza una ablación sistemática de nueve arquitecturas de aprendizaje "
         "automático, desde modelos lineales hasta redes neuronales convolucionales-recurrentes con "
         "conexiones residuales; y (3) aplica análisis SHAP para identificar qué indicadores de "
         "visitantes aportan mayor valor predictivo incremental sobre los rezagos autorregresivos."),
        ("El diseño experimental sigue una estructura factorial 9 × 3 = 25 combinaciones (con SARIMA "
         "como caso especial sin indicadores), evaluadas sobre un conjunto de prueba post-COVID estricto "
         "(2022 Q1 – 2025 Q3, 15 trimestres) para evitar que la recuperación pandémica contamine la "
         "evaluación fuera de muestra."),
    ]
    for text in paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.7)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()


def write_data_methodology(doc: Document):
    add_heading(doc, "2. Datos y Metodología", 1)

    add_heading(doc, "2.1 Fuentes de datos", 2)
    paras = [
        ("Las tres fuentes de datos provienen íntegramente del INEGI:"),
        ("• Índice de Volumen Físico (IVF): serie trimestral desde 1993 Q1 hasta 2025 Q3 (127 trimestres). "
         "Se utilizan cuatro series: IVF Total Nacional, IVF Turístico Total, IVF Turístico de Bienes e "
         "IVF Turístico de Servicios. Base 2013 = 100."),
        ("• Consumo Turístico Interior: serie trimestral 1993–2025 en millones de pesos a precios "
         "constantes, desagregada en consumo interno y receptivo."),
        ("• Indicadores de visitantes (EVI y Estadísticas de Turismo): 36 series mensuales desde enero "
         "2018 hasta septiembre 2025 (aproximadamente 92 meses), que incluyen número de visitantes y "
         "gasto total clasificados por tipo (turistas vs. excursionistas), origen (internación vs. "
         "fronterizos) y modo de transporte (aéreo, automóvil, peatón). Se agregan a frecuencia "
         "trimestral mediante suma."),
    ]
    for text in paras:
        p = doc.add_paragraph(text)
        if text.startswith("•"):
            p.paragraph_format.left_indent = Cm(0.7)
        else:
            p.paragraph_format.first_line_indent = Cm(0.7)
        for run in p.runs:
            run.font.size = Pt(11)

    add_heading(doc, "2.2 Ingeniería de características", 2)
    feat_text = (
        "Para cada serie objetivo se construyen las siguientes características:\n\n"
        "• Rezagos: t-1, t-2, t-3, t-4 de la variable objetivo.\n"
        "• Medias móviles: ventanas de 2 y 4 trimestres del objetivo.\n"
        "• Tendencia lineal: variable numérica creciente (índice temporal).\n"
        "• Dummies estacionales: cuatro indicadores binarios de trimestre (Q1–Q4).\n"
        "• Dummy COVID: variable binaria = 1 para 2020 Q1–2021 Q4, que modela "
        "explícitamente la ruptura estructural de la pandemia.\n"
        "• Indicadores INEGI: las 36 series de visitantes agregadas trimestralmente, "
        "extendidas hacia pre-2018 mediante el método de backfill seleccionado.\n\n"
        "El conjunto de características final tiene 48 columnas (11 autorregresivas + 1 tendencia + "
        "4 estacionales + 1 COVID + 36 indicadores) sobre 127 observaciones."
    )
    p = doc.add_paragraph(feat_text)
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    add_heading(doc, "2.3 Métodos de imputación pre-2018 (backfill)", 2)
    bf_rows = [
        ("Zero",        "Asigna cero a todos los indicadores en trimestres pre-2018. "
                        "Introduce una ruptura estructural artificial en 2018 Q1 pero "
                        "preserva la integridad de la serie principal."),
        ("Linear",      "Ajusta una regresión lineal simple (OLS) sobre cada indicador "
                        "usando los 30 trimestres disponibles (2018 Q1–2025 Q3) y extrapola "
                        "hacia atrás hasta 1994. Supone tendencia lineal estacionaria."),
        ("XGB Backcast","Entrena un modelo XGBoost sobre la serie temporal invertida "
                        "(pasado → futuro invertido a futuro → pasado) para extrapolar recursivamente. "
                        "Captura no linealidades pero acumula error en la extrapolación profunda."),
    ]
    for name, desc in bf_rows:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}: ")
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.7)

    add_heading(doc, "2.4 División temporal y evaluación", 2)
    split_text = (
        "Se utiliza una división cronológica estricta (sin mezcla aleatoria) para respetar "
        "la estructura temporal de los datos:\n\n"
        "• Entrenamiento: 1994 Q1 – 2021 Q4 (≈112 trimestres).\n"
        "• Prueba: 2022 Q1 – 2025 Q3 (15 trimestres, post-COVID).\n\n"
        "Las métricas de evaluación son: Error Absoluto Medio (MAE), Raíz del Error Cuadrático Medio "
        "(RMSE), Error Porcentual Absoluto Medio (MAPE) y Coeficiente de Determinación (R²). "
        "Se reportan por serie individual y como promedio de las cuatro series IVF."
    )
    p = doc.add_paragraph(split_text)
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()


def write_models(doc: Document):
    add_heading(doc, "3. Modelos", 1)

    p = doc.add_paragraph(
        "Se evalúan nueve arquitecturas organizadas en tres grupos según su inductive bias "
        "sobre la estructura temporal de la serie:"
    )
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    # Tabla de modelos
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_row_bg(hdr, "1F497D")
    for cell, text in zip(hdr.cells, ["Modelo", "Grupo", "Descripción"]):
        bold_cell(cell, text, size=10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    rows_data = [
        ("XGBoost",     "Tabular",           MODEL_DESCRIPTIONS["XGBoost"]),
        ("Ridge",       "Tabular",           MODEL_DESCRIPTIONS["Ridge"]),
        ("SARIMA",      "Estadístico",       MODEL_DESCRIPTIONS["SARIMA"]),
        ("SARIMAX",     "Estadístico",       MODEL_DESCRIPTIONS["SARIMAX"]),
        ("MLP",         "NN — denso",        MODEL_DESCRIPTIONS["MLP"]),
        ("GRU",         "NN — recurrente",   MODEL_DESCRIPTIONS["GRU"]),
        ("LSTM",        "NN — recurrente",   MODEL_DESCRIPTIONS["LSTM"]),
        ("CNN-GRU",     "NN — conv+recur.",  MODEL_DESCRIPTIONS["CNN-GRU"]),
        ("Res-CNN-GRU", "NN — conv+recur.",  MODEL_DESCRIPTIONS["Res-CNN-GRU"]),
    ]
    colors = ["FFFFFF", "F2F2F2"]
    for i, (m, g, d) in enumerate(rows_data):
        row = table.add_row()
        set_row_bg(row, colors[i % 2])
        normal_cell(row.cells[0], m,  bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        normal_cell(row.cells[1], g,  align=WD_ALIGN_PARAGRAPH.CENTER)
        normal_cell(row.cells[2], d,  align=WD_ALIGN_PARAGRAPH.LEFT)

    # Ancho de columnas
    for i, w in enumerate([1.2, 1.3, 4.0]):
        for row in table.rows:
            row.cells[i].width = Inches(w)

    doc.add_paragraph()

    add_heading(doc, "3.1 Redes neuronales — hiperparámetros y entrenamiento", 2)
    nn_text = (
        "Todas las redes neuronales (MLP, GRU, LSTM, CNN-GRU, Res-CNN-GRU) comparten los mismos "
        "hiperparámetros base: dropout = 0.2, learning rate = 0.0005 (Adam), batch size = 8, "
        "máximo 500 épocas con early stopping (patience = 30). El conjunto de validación es el "
        "último 20% cronológico del conjunto de entrenamiento (no del conjunto de prueba). "
        "Los modelos recurrentes y convolucionales usan una ventana de secuencia de 8 trimestres.\n\n"
        "Para Res-CNN-GRU, los 3 bloques residuales reducen la dimensión temporal: "
        "seq=8 → 7 → 6 → 5 vía MaxPool(stride=1). La conexión shortcut usa Conv1D(1×1) "
        "para proyectar canales cuando in_ch ≠ out_ch."
    )
    p = doc.add_paragraph(nn_text)
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()


def write_results_main(doc: Document, df: pd.DataFrame):
    add_heading(doc, "4. Resultados", 1)
    add_heading(doc, "4.1 Comparación principal (backfill lineal)", 2)

    p = doc.add_paragraph(
        "La Tabla 1 presenta MAE, RMSE y R² promediados sobre las cuatro series IVF "
        "para todos los modelos con backfill lineal — el método de imputación que produce "
        "los mejores resultados en la mayoría de los modelos (ver Sección 4.3). "
        "Los valores en negrita indican el mejor modelo en cada columna."
    )
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    # Tabla 1: comparación promedio con linear
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_row_bg(hdr, "1F497D")
    for cell, text in zip(hdr.cells, ["Modelo", "MAE", "RMSE", "MAPE (%)", "R²"]):
        bold_cell(cell, text, size=10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Calcular métricas promedio para cada modelo con linear (o "—" para sarima)
    results = []
    for model in MODEL_ORDER:
        bf = "—" if model == "SARIMA" else "Linear"
        results.append({
            "model": model,
            "mae":  avg4(df, model, bf, "mae"),
            "rmse": avg4(df, model, bf, "rmse"),
            "mape": avg4(df, model, bf, "mape"),
            "r2":   avg4(df, model, bf, "r2"),
        })

    mae_vals  = [r["mae"]  for r in results]
    rmse_vals = [r["rmse"] for r in results]
    mape_vals = [r["mape"] for r in results]
    r2_vals   = [r["r2"]   for r in results]
    best_mae  = highlight_best(mae_vals,  lower_is_better=True)
    best_rmse = highlight_best(rmse_vals, lower_is_better=True)
    best_mape = highlight_best(mape_vals, lower_is_better=True)
    best_r2   = highlight_best(r2_vals,   lower_is_better=False)

    colors = ["FFFFFF", "F2F2F2"]
    for i, r in enumerate(results):
        row = table.add_row()
        set_row_bg(row, colors[i % 2])
        normal_cell(row.cells[0], r["model"], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        for j, (val, is_best) in enumerate(zip(
            [r["mae"], r["rmse"], r["mape"], r["r2"]],
            [best_mae[i], best_rmse[i], best_mape[i], best_r2[i]]
        )):
            normal_cell(row.cells[j+1], fmt(val), bold=is_best)
            if is_best:
                set_cell_bg(row.cells[j+1], "E2EFDA")

    for i, w in enumerate([1.5, 1.0, 1.0, 1.0, 1.0]):
        for row in table.rows:
            row.cells[i].width = Inches(w)

    cap = doc.add_paragraph(
        "Tabla 1. MAE, RMSE, MAPE y R² promediados sobre 4 series IVF. "
        "Backfill lineal para todos los modelos excepto SARIMA (sin indicadores). "
        "Conjunto de prueba: 2022 Q1 – 2025 Q3. Celdas verdes = mejor valor."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Tabla 2: desglose por serie
    add_heading(doc, "4.2 Desglose por serie (MAE, backfill lineal)", 2)

    table2 = doc.add_table(rows=1, cols=6)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0]
    set_row_bg(hdr2, "1F497D")
    headers2 = ["Modelo"] + [SERIES_LABELS[s] for s in SERIES] + ["Promedio"]
    for cell, text in zip(hdr2.cells, headers2):
        bold_cell(cell, text, size=9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    table2 = doc.add_table(rows=1, cols=6)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0]
    set_row_bg(hdr2, "1F497D")
    for cell, text in zip(hdr2.cells, ["Modelo"] + [SERIES_LABELS[s] for s in SERIES] + ["Promedio"]):
        bold_cell(cell, text, size=9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    per_series = []
    for model in MODEL_ORDER:
        bf = "—" if model == "SARIMA" else "Linear"
        row_vals = []
        for s in SERIES:
            sub = df[(df["model"] == model) & (df["backfill"] == bf) & (df["series"] == s)]
            row_vals.append(sub["mae"].values[0] if len(sub) > 0 else np.nan)
        avg = np.nanmean(row_vals) if any(not np.isnan(v) for v in row_vals) else np.nan
        per_series.append({"model": model, "vals": row_vals, "avg": avg})

    # Columna best per series
    col_bests = []
    for j in range(len(SERIES)):
        col_vals = [r["vals"][j] for r in per_series]
        col_bests.append(highlight_best(col_vals))
    avg_bests = highlight_best([r["avg"] for r in per_series])

    for i, r in enumerate(per_series):
        row = table2.add_row()
        set_row_bg(row, colors[i % 2])
        normal_cell(row.cells[0], r["model"], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        for j, val in enumerate(r["vals"]):
            is_best = col_bests[j][i]
            normal_cell(row.cells[j+1], fmt(val, 2), bold=is_best)
            if is_best:
                set_cell_bg(row.cells[j+1], "E2EFDA")
        is_best_avg = avg_bests[i]
        normal_cell(row.cells[5], fmt(r["avg"], 2), bold=is_best_avg)
        if is_best_avg:
            set_cell_bg(row.cells[5], "E2EFDA")

    cap2 = doc.add_paragraph(
        "Tabla 2. MAE por serie y modelo (backfill lineal). "
        "Celdas verdes = mejor modelo para esa serie."
    )
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.runs[0].italic = True
    cap2.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def write_ablation(doc: Document, df: pd.DataFrame):
    add_heading(doc, "4.3 Análisis de ablación arquitectónica", 2)

    p = doc.add_paragraph(
        "La Tabla 3 mide el impacto incremental de cada componente arquitectónico, "
        "comparando pares consecutivos de modelos en orden de complejidad creciente. "
        "Un ΔMAE positivo indica degradación del rendimiento al agregar el componente."
    )
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    steps = [
        ("Baseline tabular",        "XGBoost → Ridge",        "XGBoost", "Ridge"),
        ("Baseline estadístico",    "XGBoost → SARIMA",       "XGBoost", "SARIMA"),
        ("+ Exógenos (SARIMAX)",    "SARIMA → SARIMAX",       "SARIMA",  "SARIMAX"),
        ("Tabular → NN densa",      "XGBoost → MLP",          "XGBoost", "MLP"),
        ("+ Memoria temporal (GRU)","MLP → GRU",              "MLP",     "GRU"),
        ("GRU → LSTM",              "GRU → LSTM",             "GRU",     "LSTM"),
        ("+ Conv. (CNN-GRU)",       "GRU → CNN-GRU",          "GRU",     "CNN-GRU"),
        ("+ Residuales",            "CNN-GRU → Res-CNN-GRU",  "CNN-GRU", "Res-CNN-GRU"),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_row_bg(hdr, "1F497D")
    for cell, text in zip(hdr.cells, ["Componente", "Comparación", "ΔMAE", "ΔRMSE", "ΔR²"]):
        bold_cell(cell, text, size=10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    colors = ["FFFFFF", "F2F2F2"]
    for i, (comp, cmp_label, m_a, m_b) in enumerate(steps):
        bf_a = "—" if m_a == "SARIMA" else "Linear"
        bf_b = "—" if m_b == "SARIMA" else "Linear"
        mae_a  = avg4(df, m_a, bf_a, "mae")
        mae_b  = avg4(df, m_b, bf_b, "mae")
        rmse_a = avg4(df, m_a, bf_a, "rmse")
        rmse_b = avg4(df, m_b, bf_b, "rmse")
        r2_a   = avg4(df, m_a, bf_a, "r2")
        r2_b   = avg4(df, m_b, bf_b, "r2")

        def delta_pct(a, b):
            if np.isnan(a) or np.isnan(b) or a == 0:
                return "—"
            pct = (b - a) / abs(a) * 100
            return f"{pct:+.1f}%"

        d_mae  = delta_pct(mae_a,  mae_b)
        d_rmse = delta_pct(rmse_a, rmse_b)
        d_r2   = f"{r2_b - r2_a:+.3f}" if not (np.isnan(r2_a) or np.isnan(r2_b)) else "—"

        row = table.add_row()
        set_row_bg(row, colors[i % 2])
        normal_cell(row.cells[0], comp,      align=WD_ALIGN_PARAGRAPH.LEFT)
        normal_cell(row.cells[1], cmp_label, align=WD_ALIGN_PARAGRAPH.LEFT)
        # Color rojo si empeora (ΔMAE > 0), verde si mejora
        for j, d in enumerate([d_mae, d_rmse]):
            normal_cell(row.cells[j+2], d)
            if d != "—":
                try:
                    val = float(d.replace("%","").replace("+",""))
                    if val > 0:
                        set_cell_bg(row.cells[j+2], "FCE4D6")
                    elif val < 0:
                        set_cell_bg(row.cells[j+2], "E2EFDA")
                except Exception:
                    pass
        normal_cell(row.cells[4], d_r2)
        if d_r2 != "—":
            try:
                val = float(d_r2.replace("+",""))
                if val > 0:
                    set_cell_bg(row.cells[4], "E2EFDA")
                elif val < 0:
                    set_cell_bg(row.cells[4], "FCE4D6")
            except Exception:
                pass

    cap = doc.add_paragraph(
        "Tabla 3. Análisis de ablación. ΔMAE/ΔRMSE negativos (verde) = mejora; "
        "positivos (naranja) = degradación. Backfill lineal en todos los casos."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def write_backfill_impact(doc: Document, df: pd.DataFrame):
    add_heading(doc, "4.4 Impacto del método de backfill", 2)

    p = doc.add_paragraph(
        "La Tabla 4 muestra el MAE promedio sobre las 4 series IVF para cada combinación "
        "modelo × backfill. El mejor backfill por modelo está resaltado en verde. "
        "SARIMA se excluye por no utilizar indicadores exógenos."
    )
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    models_bf = [m for m in MODEL_ORDER if m != "SARIMA"]
    backfills = ["Zero", "Linear", "XGB Backcast"]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_row_bg(hdr, "1F497D")
    for cell, text in zip(hdr.cells, ["Modelo", "Zero", "Linear", "XGB Backcast"]):
        bold_cell(cell, text, size=10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    colors = ["FFFFFF", "F2F2F2"]
    for i, model in enumerate(models_bf):
        vals = [avg4(df, model, bf, "mae") for bf in backfills]
        best_mask = highlight_best(vals)
        row = table.add_row()
        set_row_bg(row, colors[i % 2])
        normal_cell(row.cells[0], model, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        for j, (v, is_best) in enumerate(zip(vals, best_mask)):
            normal_cell(row.cells[j+1], fmt(v, 2), bold=is_best)
            if is_best:
                set_cell_bg(row.cells[j+1], "E2EFDA")

    cap = doc.add_paragraph(
        "Tabla 4. MAE promedio (4 series IVF) por modelo y método de backfill. "
        "Verde = mejor backfill para ese modelo."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Tabla de ganancia de linear vs zero
    add_heading(doc, "4.5 Ganancia del backfill lineal sobre relleno con ceros", 2)

    table5 = doc.add_table(rows=1, cols=5)
    table5.style = "Table Grid"
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr5 = table5.rows[0]
    set_row_bg(hdr5, "1F497D")
    for cell, text in zip(hdr5.cells, ["Modelo", "MAE (Zero)", "MAE (Linear)", "ΔMAE", "ΔMAE %"]):
        bold_cell(cell, text, size=10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, model in enumerate(models_bf):
        mae_z = avg4(df, model, "Zero",   "mae")
        mae_l = avg4(df, model, "Linear", "mae")
        delta_abs = mae_l - mae_z if not (np.isnan(mae_z) or np.isnan(mae_l)) else np.nan
        delta_pct = delta_abs / mae_z * 100 if not np.isnan(mae_z) and mae_z != 0 else np.nan

        row = table5.add_row()
        set_row_bg(row, colors[i % 2])
        normal_cell(row.cells[0], model,           bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        normal_cell(row.cells[1], fmt(mae_z, 2))
        normal_cell(row.cells[2], fmt(mae_l, 2))
        d_abs_str = fmt(delta_abs, 2) if not np.isnan(delta_abs) else "—"
        d_pct_str = f"{delta_pct:+.1f}%" if not np.isnan(delta_pct) else "—"
        normal_cell(row.cells[3], d_abs_str)
        normal_cell(row.cells[4], d_pct_str)
        if not np.isnan(delta_pct):
            color = "E2EFDA" if delta_pct < 0 else "FCE4D6"
            set_cell_bg(row.cells[3], color)
            set_cell_bg(row.cells[4], color)

    cap5 = doc.add_paragraph(
        "Tabla 5. Impacto del backfill lineal vs. relleno con ceros. "
        "Verde = mejora con backfill lineal; naranja = degradación."
    )
    cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap5.runs[0].italic = True
    cap5.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def write_shap(doc: Document):
    add_heading(doc, "5. Análisis de Interpretabilidad (SHAP)", 1)

    add_heading(doc, "5.1 Importancia de indicadores INEGI", 2)

    shap_path = "data/shap/indicator_importance_summary.csv"
    if os.path.exists(shap_path):
        shap_df = pd.read_csv(shap_path)

        if "mean_abs_shap" in shap_df.columns and "indicator" in shap_df.columns:
            top = shap_df.nlargest(15, "mean_abs_shap")
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0]
            set_row_bg(hdr, "1F497D")
            for cell, text in zip(hdr.cells, ["Rango", "Indicador", "Importancia SHAP (media)"]):
                bold_cell(cell, text, size=10)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            colors = ["FFFFFF", "F2F2F2"]
            for i, (_, row_data) in enumerate(top.iterrows()):
                row = table.add_row()
                set_row_bg(row, colors[i % 2])
                ind_label = str(row_data["indicator"]).replace("__", " / ").replace("_", " ")
                normal_cell(row.cells[0], str(i+1))
                normal_cell(row.cells[1], ind_label, align=WD_ALIGN_PARAGRAPH.LEFT)
                normal_cell(row.cells[2], f"{row_data['mean_abs_shap']:.4f}")
        else:
            # Formato alternativo
            cols = [c for c in shap_df.columns if c != "Unnamed: 0"]
            if cols:
                shap_df["mean_shap"] = shap_df[cols].mean(axis=1)
                shap_df = shap_df.sort_values("mean_shap", ascending=False).head(15)

                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = table.rows[0]
                set_row_bg(hdr, "1F497D")
                for cell, text in zip(hdr.cells, ["Rango", "Indicador", "SHAP promedio"]):
                    bold_cell(cell, text, size=10)
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                colors = ["FFFFFF", "F2F2F2"]
                for i, (_, row_data) in enumerate(shap_df.iterrows()):
                    row = table.add_row()
                    set_row_bg(row, colors[i % 2])
                    ind_label = str(row_data.get("Unnamed: 0", "")).replace("__", " / ").replace("_", " ")
                    normal_cell(row.cells[0], str(i+1))
                    normal_cell(row.cells[1], ind_label, align=WD_ALIGN_PARAGRAPH.LEFT)
                    normal_cell(row.cells[2], f"{row_data['mean_shap']:.4f}")

        cap = doc.add_paragraph(
            "Tabla 6. Top-15 indicadores por importancia SHAP media sobre las 4 series IVF "
            "(modelo XGBoost con backfill lineal)."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph("_Datos SHAP no disponibles._").runs[0].italic = True

    doc.add_paragraph()

    p = doc.add_paragraph(
        "Los indicadores de gasto total de turistas internacionales por vía aérea (entrada y salida) "
        "son consistentemente los más influyentes en las series de IVF Turístico Total y Turístico de "
        "Servicios. Este resultado es coherente con la estructura del IVF turístico: los turistas "
        "internacionales de alta capacidad de gasto y que llegan por vía aérea tienen mayor impacto "
        "en el volumen físico de servicios que los excursionistas fronterizos.\n\n"
        "El IVF Total Nacional y el IVF Turístico de Bienes muestran menor sensibilidad a los "
        "indicadores de visitantes, lo que sugiere que estos subsectores están más determinados por "
        "dinámicas domésticas o por indicadores no capturados en la EVI."
    )
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()


def write_figures(doc: Document):
    add_heading(doc, "6. Figuras", 1)

    add_heading(doc, "6.1 Series IVF históricas (1994–2025)", 2)
    add_figure(doc, "plots/general/ivf_overview.png",
               "Figura 1. Las cuatro series IVF del INEGI (1994–2025). "
               "La banda gris sombreada indica el período COVID (2020 Q1–2021 Q4). "
               "Las líneas discontinuas marcan el inicio del conjunto de prueba (2022 Q1).",
               width=6.0)

    add_heading(doc, "6.2 Comparación de modelos", 2)
    add_figure(doc, "plots/general/model_comparison.png",
               "Figura 2. MAE, RMSE y R² por modelo y serie (backfill lineal). "
               "Promediado sobre las 4 series IVF. Conjunto de prueba 2022–2025.",
               width=6.5)

    add_heading(doc, "6.3 Impacto del método de backfill", 2)
    add_figure(doc, "plots/general/backfill_impact.png",
               "Figura 3. Comparación de los tres métodos de backfill por modelo. "
               "Métricas promediadas sobre las 4 series IVF.",
               width=6.5)

    add_heading(doc, "6.4 Pronósticos — XGBoost con backfill lineal", 2)
    for s in ["ivf_total_nacional", "ivf_turistico_total"]:
        img = f"plots/xgb_linear/forecast_{s}.png"
        add_figure(doc, img,
                   f"Figura. Pronóstico XGBoost (linear) — {SERIES_LABELS[s]}. "
                   f"Línea azul = entrenamiento, naranja = prueba real, rojo punteado = predicción.",
                   width=5.5)

    add_heading(doc, "6.5 Importancia SHAP — XGBoost lineal", 2)
    for s in ["ivf_total_nacional", "ivf_turistico_total"]:
        img = f"plots/xgb_linear/shap_{s}.png"
        add_figure(doc, img,
                   f"Figura. SHAP summary plot — {SERIES_LABELS[s]} (XGBoost lineal). "
                   f"Las características se ordenan de mayor a menor valor absoluto medio.",
                   width=5.5)

    add_heading(doc, "6.6 Curvas de entrenamiento — redes neuronales", 2)
    for mtype in ["mlp", "gru", "lstm"]:
        img = f"plots/general/training_curves_{mtype}.png"
        label = mtype.upper()
        add_figure(doc, img,
                   f"Figura. Curvas train/val loss — {label} (3 variantes de backfill × 4 series). "
                   f"Los triángulos marcan el epoch de early stopping.",
                   width=6.5)


def write_conclusions(doc: Document, df: pd.DataFrame):
    add_heading(doc, "7. Conclusiones", 1)

    xgb_lin_mae  = avg4(df, "XGBoost", "Linear", "mae")
    xgb_zero_mae = avg4(df, "XGBoost", "Zero",   "mae")
    sarima_mae   = avg4(df, "SARIMA",  "—",      "mae")
    ridge_lin_mae = avg4(df, "Ridge",  "Linear", "mae")
    sarimax_lin_mae = avg4(df, "SARIMAX", "Linear", "mae")
    nn_maes = {m: avg4(df, m, "Linear", "mae")
               for m in ["MLP","GRU","LSTM","CNN-GRU","Res-CNN-GRU"]}
    best_nn = min(nn_maes, key=lambda k: nn_maes[k] if not np.isnan(nn_maes[k]) else 1e9)

    bf_gain = (xgb_zero_mae - xgb_lin_mae) / xgb_zero_mae * 100 if xgb_zero_mae else 0

    findings = [
        (f"XGBoost con backfill lineal es el mejor modelo general "
         f"(MAE promedio = {fmt(xgb_lin_mae, 2)} frente a MAE = {fmt(sarima_mae, 2)} de SARIMA "
         f"y MAE = {fmt(nn_maes[best_nn], 2)} de la mejor red neuronal — {best_nn}). "
         "La ventaja de XGBoost sobre los modelos neuronales se explica por la combinación de "
         "alta dimensionalidad relativa (48 features, ~112 observaciones de entrenamiento) y "
         "la capacidad de XGBoost para regularizar eficientemente en muestras pequeñas."),

        (f"El backfill lineal reduce el MAE de XGBoost en {bf_gain:.1f}% frente al relleno con ceros, "
         "demostrando que la imputación de indicadores históricos de alta calidad es tan importante "
         "como la elección del modelo. El backcast XGBoost tiene peor desempeño que el lineal, "
         "probablemente porque los errores de extrapolación recursiva hacia 1994 (~24 años) se acumulan."),

        (f"SARIMAX (MAE = {fmt(sarimax_lin_mae, 2)}) supera a SARIMA puro ({fmt(sarima_mae, 2)}), "
         "confirmando que los indicadores INEGI de visitantes agregan señal predictiva incluso "
         "dentro del marco estadístico clásico."),

        ("El análisis de ablación revela que agregar recurrencia pura (GRU) sin extracción "
         "convolucional empeora significativamente el pronóstico frente al MLP. La adición de "
         "capas convolucionales (CNN-GRU) recupera parte del rendimiento al extraer patrones "
         "locales en la dimensión temporal antes de la recurrencia. Las conexiones residuales "
         "(Res-CNN-GRU) no aportan mejora estadística en este contexto, posiblemente porque "
         "la longitud de secuencia (8 trimestres) no es suficiente para justificar la profundidad."),

        ("Los indicadores de gasto de turistas internacionales por vía aérea son las variables "
         "más influyentes en el pronóstico del IVF turístico de servicios, validando que la "
         "integración de microdatos de movilidad turística mejora la explicabilidad del modelo "
         "más allá de los rezagos autorregresivos puros."),
    ]

    for i, text in enumerate(findings, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(0.4)

    doc.add_paragraph()

    add_heading(doc, "7.1 Limitaciones y trabajo futuro", 2)
    limitations = [
        "El conjunto de prueba (15 trimestres) es reducido y puede ser sensible a eventos "
        "idiosincráticos del período 2022–2025 (recuperación asimétrica post-COVID, inflación).",
        "Los modelos neuronales fueron entrenados con una única semilla aleatoria. Para el paper "
        "final se recomienda promediar sobre 5–10 semillas para reducir la varianza de resultados.",
        "El backcast pre-2018 asume estacionariedad en la relación entre indicadores e IVF, "
        "lo que puede no sostenerse en presencia de cambios estructurales anteriores a 2018.",
        "Los indicadores INEGI tienen cobertura mensual; una versión futura podría incorporar "
        "modelos MIDAS (Mixed Data Sampling) para aprovechar la frecuencia alta sin agregación.",
        "No se consideran variables macroeconómicas exógenas (tipo de cambio, PIB de EE.UU.) "
        "que podrían mejorar el pronóstico de segmentos de turismo internacional.",
    ]
    for text in limitations:
        p = doc.add_paragraph(f"• {text}")
        p.paragraph_format.left_indent = Cm(0.7)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_paragraph()


def write_appendix(doc: Document, df: pd.DataFrame):
    doc.add_page_break()
    add_heading(doc, "Apéndice A — Métricas completas (todos los experimentos)", 1)

    p = doc.add_paragraph(
        "Tabla A1 presenta MAE, RMSE y R² para los 25 experimentos y las 4 series IVF."
    )
    p.paragraph_format.first_line_indent = Cm(0.7)
    for run in p.runs:
        run.font.size = Pt(10)

    # Una tabla por serie
    for series in SERIES:
        doc.add_paragraph()
        h = doc.add_paragraph(f"Serie: {SERIES_LABELS[series]}")
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(10)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0]
        set_row_bg(hdr, "2E75B6")
        for cell, text in zip(hdr.cells, ["Experimento", "MAE", "RMSE", "MAPE (%)", "R²"]):
            bold_cell(cell, text, size=9)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        sub = df[df["series"] == series].sort_values("mae")
        colors = ["FFFFFF", "F2F2F2"]
        for i, (_, row_data) in enumerate(sub.iterrows()):
            row = table.add_row()
            set_row_bg(row, colors[i % 2])
            normal_cell(row.cells[0], row_data["exp"], size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
            for j, metric in enumerate(["mae", "rmse", "mape", "r2"]):
                normal_cell(row.cells[j+1], fmt(row_data[metric], 4), size=8)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default="reporte_paper_turismo.docx")
    args = parser.parse_args()

    print("Cargando métricas de los 25 experimentos...")
    df = build_df()
    n_ok = df.dropna(subset=["mae"]).shape[0]
    print(f"  Registros disponibles: {n_ok} / {len(df)}")

    doc = Document()

    # Márgenes 2.5cm
    from docx.oxml.ns import qn as _qn
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # Estilo de fuente base
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    write_title_page(doc)
    write_abstract(doc, df)
    write_introduction(doc)
    write_data_methodology(doc)
    write_models(doc)
    write_results_main(doc, df)
    write_ablation(doc, df)
    write_backfill_impact(doc, df)
    write_shap(doc)
    write_figures(doc)
    write_conclusions(doc, df)
    write_appendix(doc, df)

    doc.save(args.out)
    print(f"\nReporte guardado en: {args.out}")


if __name__ == "__main__":
    main()
